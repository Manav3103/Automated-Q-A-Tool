import PyPDF2
from docx import Document
import os

class DocumentProcessor:
    """Handles text extraction from PDF, DOCX, and TXT files"""
    
    def extract_text(self, file_path):
        """Extract text from PDF, DOCX, or TXT file"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return self._extract_from_docx(file_path)
        elif file_extension == '.txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file using PyPDF2"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
        
        return text.strip()
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX file using python-docx"""
        text = ""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
                    
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
        
        return text.strip()
    
    def _extract_from_txt(self, file_path):
        """Extract text from TXT file"""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            except Exception as e:
                raise Exception(f"Error reading TXT file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")
        
        return text.strip()
